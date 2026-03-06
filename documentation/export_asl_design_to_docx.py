#!/usr/bin/env python3
"""
One-off script to export ASL_INTEGRATION_TECHNICAL_DESIGN.md to a Word document
with Times New Roman font, for uploading to Google Docs.
Requires: pip install python-docx
Run from repo root: python documentation/export_asl_design_to_docx.py
"""
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    raise SystemExit("Please install python-docx: pip install python-docx")

REPO_ROOT = Path(__file__).resolve().parent.parent
MD_PATH = REPO_ROOT / "documentation" / "ASL_INTEGRATION_TECHNICAL_DESIGN.md"
OUT_PATH = REPO_ROOT / "documentation" / "ASL_INTEGRATION_TECHNICAL_DESIGN.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = 12
FONT_SIZE_H1 = 16
FONT_SIZE_H2 = 14
FONT_SIZE_H3 = 12


def set_run_font(run, font_name=FONT_NAME, size_pt=None, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt:
        run.font.size = Pt(size_pt)
    run.font.bold = bold


def set_paragraph_font(paragraph, font_name=FONT_NAME, size_pt=FONT_SIZE_BODY, bold=False):
    for run in paragraph.runs:
        set_run_font(run, font_name=font_name, size_pt=size_pt, bold=bold)


def set_style_font(style, font_name=FONT_NAME, size_pt=FONT_SIZE_BODY):
    style.font.name = font_name
    if style.element.rPr is None:
        style.element.get_or_add_rPr()
    try:
        style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass
    style.font.size = Pt(size_pt)


def add_paragraph(doc, text, style="Normal", font_size=FONT_SIZE_BODY):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, size_pt=font_size)
    return p


def add_heading(doc, text, level=1):
    if level == 1:
        p = doc.add_heading(text, level=1)
        for run in p.runs:
            set_run_font(run, size_pt=FONT_SIZE_H1, bold=True)
    elif level == 2:
        p = doc.add_heading(text, level=2)
        for run in p.runs:
            set_run_font(run, size_pt=FONT_SIZE_H2, bold=True)
    else:
        p = doc.add_heading(text, level=3)
        for run in p.runs:
            set_run_font(run, size_pt=FONT_SIZE_H3, bold=True)
    return p


def add_table(doc, rows, font_size=11):
    cols = len(rows[0]) if rows else 0
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = str(cell_text).strip()
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size_pt=font_size)
    return table


def add_code_block(doc, text, font_size=10):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(font_size)
    return p


def main():
    doc = Document()
    # Set default font for Normal style
    style = doc.styles["Normal"]
    set_style_font(style, size_pt=FONT_SIZE_BODY)
    for heading_level in range(1, 4):
        h_style = doc.styles[f"Heading {heading_level}"]
        set_style_font(h_style, size_pt=FONT_SIZE_H1 if heading_level == 1 else FONT_SIZE_H2 if heading_level == 2 else FONT_SIZE_H3)

    content = MD_PATH.read_text(encoding="utf-8")
    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    table_sep = re.compile(r"^\|?[-:\s|]+\|?$")

    while i < len(lines):
        line = lines[i]
        raw = line

        # Code block
        if line.strip().startswith("```"):
            if in_code_block:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table: detect table rows
        if "|" in line and line.strip().startswith("|"):
            if table_sep.match(line.strip()):
                i += 1
                continue  # skip separator row
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if cells:
                table_rows.append(cells)
            in_table = True
            i += 1
            continue
        if in_table and table_rows:
            add_table(doc, table_rows)
            table_rows = []
            in_table = False
            # don't increment; reprocess current line

        # Headings
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), level=1)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        # List items
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:].strip()
            # Strip markdown bold
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(text)
            set_run_font(run, size_pt=FONT_SIZE_BODY)
            i += 1
            continue
        if re.match(r"^\d+\.\s", line.strip()):
            text = re.sub(r"^\d+\.\s", "", line.strip())
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(text)
            set_run_font(run, size_pt=FONT_SIZE_BODY)
            i += 1
            continue

        # Empty line
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Ordinary paragraph: strip markdown links and bold for simplicity
        text = line.strip()
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
        text = re.sub(r"`([^`]+)`", r"\1", text)
        add_paragraph(doc, text)
        i += 1
        continue

    if table_rows:
        add_table(doc, table_rows)

    # Apply Times New Roman to all paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.name != "Consolas":
                run.font.name = FONT_NAME
                try:
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
                except Exception:
                    pass
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = FONT_NAME

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")
    print("Upload this file to Google Drive, then right-click → Open with → Google Docs.")


if __name__ == "__main__":
    main()
